from ..constants import OPT_ENGINE, OPT_REFERENCE, OPT_DOCX
from ..utils import WHITE, BLACK
from ..database.logic import load_analysis, get_analysis_engine_id, MOVE_INDEX, SAN_INDEX, ANNOTATION_INDEX, \
    EVALUATION_INDEX, CP_LOSS_INDEX, WIN_PERCENT_INDEX, ACCURACY_INDEX
from .constants import ANALYSIS_HEADERS, SUMMARY_HEADERS
from ..analysis.calculations import calculate_summary_statistics, extract_player_analysis
from .images import export_board_image_after_halfmoves, export_win_percent_chart_image
from .game_info import load_game_information
from docx import Document
from docx.shared import Pt
from pathlib import Path

REPORT_FONT_SIZE = 9


def get_analysis_table_headers_for_report_document():
    """
    Return the headers for the analysis tables that are written to the analysis report

    :return: List of column headers
    """
    return [
        ANALYSIS_HEADERS[MOVE_INDEX],
        ANALYSIS_HEADERS[SAN_INDEX],
        ANALYSIS_HEADERS[ANNOTATION_INDEX],
        ANALYSIS_HEADERS[EVALUATION_INDEX],
        ANALYSIS_HEADERS[CP_LOSS_INDEX],
        ANALYSIS_HEADERS[WIN_PERCENT_INDEX],
        ANALYSIS_HEADERS[ACCURACY_INDEX]
    ]


def get_player_analysis_for_report_document(analysis, player):
    """
    Return the analysis data for the specified player, including only those columns that are
    written to the analysis report

    :param analysis: Detailed per-move analysis for the whole game
    :param player: Player to extract the data for
    :return: Analysis data for the report for the specifed player
    """
    player_analysis = extract_player_analysis(analysis, player)
    return [
        [
            x[MOVE_INDEX],
            x[SAN_INDEX],
            x[ANNOTATION_INDEX],
            x[EVALUATION_INDEX],
            x[CP_LOSS_INDEX],
            f"{x[WIN_PERCENT_INDEX]:.2f}",
            f"{x[ACCURACY_INDEX]:.2f}"
        ]
        for x in player_analysis
    ]


def add_table_to_analysis_document(document, headers, table_data):
    """
    Add a table to the analysis report

    :param document: Document to add the table to
    :param headers: List of table headers
    :param table_data: List of lists of data, one per row, to populate the table
    """

    # Calculate the table size
    rows = len(table_data)
    columns = len(headers)

    # Create the table
    table = document.add_table(rows=rows + 1, cols=columns)

    # Populate the headers
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Populate the rows
    for i, data in enumerate(table_data):
        for j, value in enumerate(data):
            table.rows[1 + i].cells[j].text = f"{value:.2f}" if isinstance(value, float) else str(value)

    # Style and autofit the table
    table.autofit = True
    table.style = "Light Shading Accent 1"


def export_analysis_document(options):
    """
    Generate an analysis report document

    :param options: Dictionary of reporting parameters
    """

    # Load the analysis results
    analysis_engine_id = get_analysis_engine_id(options[OPT_ENGINE])
    analysis = load_analysis(options[OPT_REFERENCE], analysis_engine_id)

    # Calculate summary statistics
    summary_statistics = calculate_summary_statistics(analysis)

    # Generate an image of the final position
    board_position_image = export_board_image_after_halfmoves(options[OPT_REFERENCE], "*", None)

    # Generate an image of the Win Chance Chart
    win_percent_image = export_win_percent_chart_image(analysis)

    # Create the document
    document = Document()

    # Set the font size for normal text
    style = document.styles["Normal"]
    font = style.font
    font.size = Pt(REPORT_FONT_SIZE)

    # Add the document title
    document.add_heading(f"Analysis of Game '{options[OPT_REFERENCE]}'", 0)

    # Add the image of the board
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(board_position_image)

    # Add the game information
    info = load_game_information(options[OPT_REFERENCE], False, options[OPT_ENGINE], True)
    if info:
        document.add_heading(f"Game Information", level=1)
        add_table_to_analysis_document(document, ["Item", "Value"], info)

    # Move to the next page
    document.add_page_break()

    # Add the summary statistics table
    document.add_heading(f"Analysis Summary", level=1)
    add_table_to_analysis_document(document, SUMMARY_HEADERS, summary_statistics)

    # Add the win chance chart
    document.add_heading(f"Win % Chart", level=1)
    document.add_paragraph("")
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(win_percent_image)

    # Add the two analysis tables, one for each player
    headers = get_analysis_table_headers_for_report_document()
    for player in [WHITE, BLACK]:
        player_analysis = get_player_analysis_for_report_document(analysis, player)
        document.add_heading(f"Analysis for {player}", level=1)
        add_table_to_analysis_document(document, headers, player_analysis)

    # Save the report
    document.save(options[OPT_DOCX])

    # Remove the intermediate image files
    Path(board_position_image).resolve().unlink()
    Path(win_percent_image).resolve().unlink()
